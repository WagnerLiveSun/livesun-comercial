"""
Blueprint de Contratos de Prestação de Serviços - Rotas e Lógica de Negócio
"""

from flask import Blueprint, render_template, request, jsonify, redirect, url_for, flash, current_app, send_file, make_response
from flask_login import login_required, current_user
from sqlalchemy import and_, or_, func, desc
from datetime import datetime, date, timedelta
from decimal import Decimal
import json

from src.models import db, Entidade, Orcamento
from src.models.contratos import (
    ClausulaContratoPadrao,
    Contrato,
    ContratoClausula,
    ContratoHistorico,
    ContratoAnexo,
    ContratoParametro,
    ContratoParametroValor
)
from src.tenant import tenant_id, scoped_query, scoped_get_or_404, validate_ownership
from src.access_control import require_permission

contratos_bp = Blueprint('contratos', __name__, url_prefix='/contratos', template_folder='../templates/contratos')


# ============================================================================
# UTILITÁRIOS E HELPERS
# ============================================================================

def _gerar_numero_sequencial(prefixo, empresa_id=None):
    """Gera número sequencial para contratos"""
    if not empresa_id:
        empresa_id = tenant_id()
    
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    return f"{prefixo}-{timestamp}"


def _criar_historico(contrato_id, acao, descricao='', campos_alterados=None, clausulas_alteradas=None, snapshot=None):
    """Cria registro de histórico do contrato"""
    try:
        # Obter última versão
        ultimo_historico = ContratoHistorico.query.filter_by(contrato_id=contrato_id).order_by(ContratoHistorico.versao.desc()).first()
        nova_versao = (ultimo_historico.versao + 1) if ultimo_historico else 1
        
        contrato = Contrato.query.get(contrato_id)
        
        historico = ContratoHistorico(
            contrato_id=contrato_id,
            versao=nova_versao,
            acao=acao,
            status_anterior=contrato.status if contrato else None,
            status_novo=contrato.status if contrato else None,
            descricao_alteracao=descricao,
            campos_alterados=campos_alterados,
            clausulas_alteradas=clausulas_alteradas,
            alterado_por_user_id=current_user.id if current_user.is_authenticated else None,
            data_alteracao=datetime.now(),
            snapshot_contrato=snapshot
        )
        
        db.session.add(historico)
        db.session.commit()
        
        return historico
    except Exception as e:
        current_app.logger.error(f"Erro ao criar histórico: {str(e)}")
        db.session.rollback()
        return None


def _substituir_placeholders(texto, contrato):
    """Substitui placeholders no texto pelos valores do contrato"""
    if not texto:
        return texto
    
    # Dados da empresa (contratada)
    empresa = contrato.empresa
    contratada = contrato.contratada
    contratante = contrato.contratante
    
    # Mapeamento de placeholders
    placeholders = {
        '{CONTRATADA_RAZAO_SOCIAL}': contratada.nome if contratada else '',
        '{CONTRATADA_CNPJ}': contratada.cnpj_cpf if contratada else '',
        '{CONTRATADA_ENDERECO}': f"{contratada.endereco_rua or ''}, {contratada.endereco_numero or ''}" if contratada else '',
        '{CONTRATADA_CIDADE}': contratada.endereco_cidade or '' if contratada else '',
        '{CONTRATADA_UF}': contratada.endereco_uf or '' if contratada else '',
        '{CONTRATANTE_RAZAO_SOCIAL}': contratante.nome if contratante else '',
        '{CONTRATANTE_CNPJ_CPF}': contratante.cnpj_cpf if contratante else '',
        '{CONTRATANTE_ENDERECO}': f"{contratante.endereco_rua or ''}, {contratante.endereco_numero or ''}" if contratante else '',
        '{CONTRATANTE_CIDADE}': contratante.endereco_cidade or '' if contratante else '',
        '{CONTRATANTE_UF}': contratante.endereco_uf or '' if contratante else '',
        '{CONTRATO_VALOR_TOTAL}': f"R$ {contrato.valor_total:,.2f}" if contrato.valor_total else '',
        '{CONTRATO_VALOR_MENSAL}': f"R$ {contrato.valor_mensal:,.2f}" if contrato.valor_mensal else '',
        '{CONTRATO_FORMA_PAGAMENTO}': contrato.forma_pagamento or '',
        '{CONTRATO_PERIODICIDADE}': contrato.periodicidade or '',
        '{CONTRATO_DATA_INICIO}': contrato.data_inicio_vigencia.strftime('%d/%m/%Y') if contrato.data_inicio_vigencia else '',
        '{CONTRATO_DATA_FIM}': contrato.data_fim_vigencia.strftime('%d/%m/%Y') if contrato.data_fim_vigencia else '',
        '{CONTRATO_DESCRICAO_SERVICOS}': contrato.descricao_servicos or '',
        '{CONTRATO_NUMERO}': contrato.numero,
        '{CONTRATO_DATA_ASSINATURA}': contrato.data_assinatura.strftime('%d/%m/%Y') if contrato.data_assinatura else '',
    }
    
    # Substituir placeholders
    for placeholder, valor in placeholders.items():
        texto = texto.replace(placeholder, valor)
    
    return texto


# ============================================================================
# 1. CLÁUSULAS PADRÃO
# ============================================================================

@contratos_bp.route('/clausulas', methods=['GET'])
@login_required
@require_permission('contratos_view')
def clausulas_lista():
    """Lista cláusulas padrão"""
    empresa_id = tenant_id()
    
    tipo = request.args.get('tipo', '')
    categoria = request.args.get('categoria', '')
    busca = request.args.get('busca', '')
    
    query = ClausulaContratoPadrao.query.filter_by(empresa_id=empresa_id, ativo=True)
    
    if tipo:
        query = query.filter_by(tipo=tipo)
    if categoria:
        query = query.filter_by(categoria=categoria)
    if busca:
        query = query.filter(
            (ClausulaContratoPadrao.titulo.ilike(f'%{busca}%')) |
            (ClausulaContratoPadrao.codigo.ilike(f'%{busca}%'))
        )
    
    clausulas = query.order_by(ClausulaContratoPadrao.ordem_padrao, ClausulaContratoPadrao.titulo).all()
    
    return render_template('clausulas/lista.html', clausulas=clausulas, tipo_filtro=tipo, categoria_filtro=categoria, busca=busca)


@contratos_bp.route('/clausulas/nova', methods=['GET', 'POST'])
@login_required
@require_permission('contratos_create')
def clausulas_criar():
    """Cria nova cláusula padrão"""
    empresa_id = tenant_id()
    
    if request.method == 'POST':
        try:
            clausula = ClausulaContratoPadrao(
                empresa_id=empresa_id,
                codigo=request.form.get('codigo'),
                titulo=request.form.get('titulo'),
                texto_base=request.form.get('texto_base'),
                descricao=request.form.get('descricao'),
                tipo=request.form.get('tipo', 'opcional'),
                editavel=request.form.get('editavel') == 'on',
                ordem_padrao=int(request.form.get('ordem_padrao', 0)),
                categoria=request.form.get('categoria'),
                tipo_contrato=request.form.get('tipo_contrato'),
                ativo=True,
                criado_por_user_id=current_user.id
            )
            
            db.session.add(clausula)
            db.session.commit()
            
            flash(f'Cláusula {clausula.titulo} criada com sucesso!', 'success')
            return redirect(url_for('contratos.clausulas_lista'))
        
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao criar cláusula: {str(e)}', 'danger')
    
    return render_template('clausulas/formulario.html', clausula=None, action='criar')


@contratos_bp.route('/clausulas/<int:id>/editar', methods=['GET', 'POST'])
@login_required
@require_permission('contratos_edit')
def clausulas_editar(id):
    """Edita cláusula padrão"""
    clausula = scoped_get_or_404(ClausulaContratoPadrao, id)
    
    if request.method == 'POST':
        try:
            dados_anteriores = {
                'titulo': clausula.titulo,
                'texto_base': clausula.texto_base[:100] + '...' if len(clausula.texto_base) > 100 else clausula.texto_base
            }
            
            clausula.codigo = request.form.get('codigo')
            clausula.titulo = request.form.get('titulo')
            clausula.texto_base = request.form.get('texto_base')
            clausula.descricao = request.form.get('descricao')
            clausula.tipo = request.form.get('tipo', 'opcional')
            clausula.editavel = request.form.get('editavel') == 'on'
            clausula.ordem_padrao = int(request.form.get('ordem_padrao', 0))
            clausula.categoria = request.form.get('categoria')
            clausula.tipo_contrato = request.form.get('tipo_contrato')
            clausula.ativo = request.form.get('ativo') == 'on'
            clausula.atualizado_por_user_id = current_user.id
            
            db.session.commit()
            
            flash(f'Cláusula {clausula.titulo} atualizada com sucesso!', 'success')
            return redirect(url_for('contratos.clausulas_lista'))
        
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao atualizar cláusula: {str(e)}', 'danger')
    
    return render_template('clausulas/formulario.html', clausula=clausula, action='editar')


# ============================================================================
# 2. CONTRATOS
# ============================================================================

@contratos_bp.route('/', methods=['GET'])
@login_required
@require_permission('contratos_view')
def index():
    """Lista contratos"""
    empresa_id = tenant_id()
    
    status = request.args.get('status', '')
    cliente_id = request.args.get('cliente_id', '')
    
    query = Contrato.query.filter_by(empresa_id=empresa_id)
    
    if status:
        query = query.filter_by(status=status)
    if cliente_id:
        query = query.filter_by(cliente_id=int(cliente_id))
    
    contratos = query.order_by(desc(Contrato.data_geracao)).all()
    
    clientes = Entidade.query.filter_by(empresa_id=empresa_id, tipo='C', ativo=True).order_by(Entidade.nome).all()
    
    return render_template('index.html', contratos=contratos, clientes=clientes, status_filtro=status, cliente_id_filtro=cliente_id)


@contratos_bp.route('/<int:id>/reordenar-clausulas', methods=['POST'])
@login_required
@require_permission('contratos_edit')
def reordenar_clausulas(id):
    """Reordena as cláusulas de um contrato"""
    contrato = scoped_get_or_404(Contrato, id)

    if contrato.empresa_id != tenant_id():
        return jsonify({'error': 'Não autorizado'}), 403

    try:
        data = request.get_json()
        clausula_ids = data.get('clausula_ids', [])

        # Atualizar ordem das cláusulas
        for index, clausula_id in enumerate(clausula_ids):
            clausula = ContratoClausula.query.filter_by(
                id=clausula_id,
                contrato_id=contrato.id
            ).first()
            if clausula:
                clausula.ordem = index + 1

        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@contratos_bp.route('/gerar-orcamento/<int:orcamento_id>', methods=['POST'])
@login_required
@require_permission('contratos_create')
def gerar_orcamento(orcamento_id):
    """Gera contrato a partir de orçamento aprovado"""
    orcamento = scoped_get_or_404(Orcamento, orcamento_id)
    
    if orcamento.status != 'aprovado':
        flash('Apenas orçamentos aprovados podem gerar contratos.', 'danger')
        return redirect(url_for('comercial_operacional.orcamentos_detalhe', orcamento_id=orcamento_id))
    
    try:
        empresa_id = tenant_id()
        
        # Buscar entidade da empresa (contratada)
        empresa = current_user.empresa
        contratada = Entidade.query.filter_by(empresa_id=empresa_id, tipo='F').first()
        
        if not contratada:
            flash('Não foi encontrada uma entidade de fornecedor para a empresa.', 'danger')
            return redirect(url_for('comercial_operacional.orcamentos_detalhe', orcamento_id=orcamento_id))
        
        # Criar contrato
        contrato = Contrato(
            empresa_id=empresa_id,
            numero=_gerar_numero_sequencial('CTR'),
            serie='CTR',
            titulo=f"Contrato de Prestação de Serviços - {orcamento.cliente.nome}",
            orcamento_id=orcamento.id,
            cliente_id=orcamento.cliente_id,
            vendedor_id=orcamento.vendedor_id,
            contratada_entidade_id=contratada.id,
            contratante_entidade_id=orcamento.cliente_id,
            valor_total=Decimal(str(orcamento.valor_total)) if orcamento.valor_total else Decimal('0.00'),
            valor_mensal=None,  # Será definido manualmente no contrato
            forma_pagamento=None,  # Será definido manualmente no contrato
            periodicidade=None,  # Será definido manualmente no contrato
            data_inicio_vigencia=date.today(),
            data_fim_vigencia=None,
            status='rascunho',
            descricao_servicos=orcamento.observacoes or '',
            objeto_contrato=orcamento.observacoes or '',
            data_geracao=datetime.now(),
            gerado_por_user_id=current_user.id
        )
        
        db.session.add(contrato)
        db.session.flush()
        
        # Gerar cláusulas automáticas
        clausulas_padrao = ClausulaContratoPadrao.query.filter_by(
            empresa_id=empresa_id,
            tipo='obrigatoria',
            ativo=True
        ).order_by(ClausulaContratoPadrao.ordem_padrao).all()
        
        for idx, clausula_padrao in enumerate(clausulas_padrao, 1):
            clausula = ContratoClausula(
                contrato_id=contrato.id,
                clausula_padrao_id=clausula_padrao.id,
                titulo=clausula_padrao.titulo,
                texto=clausula_padrao.texto_base,
                ordem=idx,
                editavel=clausula_padrao.editavel,
                obrigatoria=True
            )
            db.session.add(clausula)
        
        # Criar histórico
        _criar_historico(contrato.id, 'criado', f'Contrato gerado a partir do orçamento {orcamento.numero}')
        
        db.session.commit()
        
        flash(f'Contrato {contrato.numero} gerado com sucesso!', 'success')
        return redirect(url_for('contratos.editar', id=contrato.id))
    
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Erro ao gerar contrato: {str(e)}")
        flash(f'Erro ao gerar contrato: {str(e)}', 'danger')
        return redirect(url_for('comercial_operacional.orcamentos_detalhe', orcamento_id=orcamento_id))


@contratos_bp.route('/<int:id>/editar', methods=['GET', 'POST'])
@login_required
@require_permission('contratos_edit')
def editar(id):
    """Edita contrato (minuta)"""
    contrato = scoped_get_or_404(Contrato, id)

    if contrato.empresa_id != tenant_id():
        flash('Não autorizado', 'danger')
        return redirect(url_for('contratos.index'))

    # Carregar parâmetros disponíveis
    parametros = ContratoParametro.query.filter_by(ativo=True).order_by(ContratoParametro.nome).all()

    # Criar dicionário de valores de placeholders
    parametros_valores_dict = {}
    for pv in contrato.parametros_valores:
        parametros_valores_dict[pv.parametro.codigo] = pv.valor

    # Adicionar ao contrato para uso no template
    contrato.parametros_valores_dict = parametros_valores_dict

    if request.method == 'POST':
        acao = request.form.get('acao', '')

        if acao == 'salvar_placeholders':
            try:
                for parametro in parametros:
                    valor = request.form.get(f'placeholder_{parametro.codigo}', '')
                    if valor:
                        # Verificar se já existe valor para este parâmetro
                        pv = ContratoParametroValor.query.filter_by(
                            contrato_id=contrato.id,
                            parametro_id=parametro.id
                        ).first()

                        if pv:
                            pv.valor = valor
                        else:
                            pv = ContratoParametroValor(
                                contrato_id=contrato.id,
                                parametro_id=parametro.id,
                                valor=valor
                            )
                            db.session.add(pv)

                db.session.commit()
                flash('Placeholders salvos com sucesso!', 'success')
            except Exception as e:
                db.session.rollback()
                flash(f'Erro ao salvar placeholders: {str(e)}', 'danger')

            return redirect(url_for('contratos.editar', id=id))

        elif acao == 'salvar_dados':
            try:
                dados_anteriores = {
                    'valor_total': str(contrato.valor_total),
                    'valor_mensal': str(contrato.valor_mensal) if contrato.valor_mensal else None,
                    'forma_pagamento': contrato.forma_pagamento,
                    'periodicidade': contrato.periodicidade
                }
                
                contrato.titulo = request.form.get('titulo')
                contrato.valor_total = Decimal(request.form.get('valor_total', 0))
                contrato.valor_mensal = Decimal(request.form.get('valor_mensal')) if request.form.get('valor_mensal') else None
                contrato.forma_pagamento = request.form.get('forma_pagamento')
                contrato.periodicidade = request.form.get('periodicidade')
                contrato.data_inicio_vigencia = datetime.strptime(request.form.get('data_inicio_vigencia'), '%Y-%m-%d').date() if request.form.get('data_inicio_vigencia') else None
                contrato.data_fim_vigencia = datetime.strptime(request.form.get('data_fim_vigencia'), '%Y-%m-%d').date() if request.form.get('data_fim_vigencia') else None
                contrato.descricao_servicos = request.form.get('descricao_servicos')
                contrato.objeto_contrato = request.form.get('objeto_contrato')
                
                db.session.commit()
                
                dados_novos = {
                    'valor_total': str(contrato.valor_total),
                    'valor_mensal': str(contrato.valor_mensal) if contrato.valor_mensal else None,
                    'forma_pagamento': contrato.forma_pagamento,
                    'periodicidade': contrato.periodicidade
                }
                
                _criar_historico(contrato.id, 'editado', 'Dados do contrato atualizados', dados_anteriores, dados_novos)
                
                flash('Dados do contrato atualizados!', 'success')
            
            except Exception as e:
                db.session.rollback()
                flash(f'Erro ao atualizar contrato: {str(e)}', 'danger')
        
        elif acao == 'adicionar_clausula':
            try:
                clausula_padrao_id = request.form.get('clausula_padrao_id')
                if clausula_padrao_id:
                    clausula_padrao = ClausulaContratoPadrao.query.get(int(clausula_padrao_id))
                    
                    # Obter próxima ordem
                    ultima_ordem = db.session.query(func.max(ContratoClausula.ordem)).filter_by(contrato_id=contrato.id).scalar() or 0
                    
                    clausula = ContratoClausula(
                        contrato_id=contrato.id,
                        clausula_padrao_id=clausula_padrao.id,
                        titulo=clausula_padrao.titulo,
                        texto=clausula_padrao.texto_base,
                        ordem=ultima_ordem + 1,
                        editavel=clausula_padrao.editavel,
                        obrigatoria=False
                    )
                    
                    db.session.add(clausula)
                    db.session.commit()
                    
                    flash('Cláusula adicionada!', 'success')
            
            except Exception as e:
                db.session.rollback()
                flash(f'Erro ao adicionar cláusula: {str(e)}', 'danger')
        
        elif acao == 'remover_clausula':
            try:
                clausula_id = request.form.get('clausula_id')
                clausula = ContratoClausula.query.filter_by(id=int(clausula_id), contrato_id=contrato.id).first_or_404()

                if clausula.obrigatoria:
                    flash('Não é possível remover cláusulas obrigatórias.', 'warning')
                else:
                    db.session.delete(clausula)
                    db.session.commit()
                    flash('Cláusula removida!', 'success')

            except Exception as e:
                db.session.rollback()
                flash(f'Erro ao remover cláusula: {str(e)}', 'danger')

        elif acao == 'adicionar_anexo':
            try:
                arquivo = request.files.get('arquivo')
                if not arquivo:
                    flash('Selecione um arquivo para anexar.', 'warning')
                    return redirect(url_for('contratos.editar', id=id))

                # Salvar arquivo
                import os
                import uuid
                from werkzeug.utils import secure_filename

                # Criar diretório de uploads se não existir
                upload_dir = os.path.join(current_app.root_path, 'uploads', 'contratos')
                os.makedirs(upload_dir, exist_ok=True)

                # Gerar nome único
                filename = secure_filename(arquivo.filename)
                unique_filename = f"{uuid.uuid4()}_{filename}"
                filepath = os.path.join(upload_dir, unique_filename)

                arquivo.save(filepath)

                # Criar registro de anexo
                anexo = ContratoAnexo(
                    contrato_id=contrato.id,
                    nome_arquivo=filename,
                    caminho_arquivo=filepath,
                    tipo_mime=arquivo.content_type,
                    tamanho_bytes=os.path.getsize(filepath),
                    upload_por_user_id=current_user.id
                )
                db.session.add(anexo)
                db.session.commit()

                flash('Anexo adicionado com sucesso!', 'success')
            except Exception as e:
                db.session.rollback()
                flash(f'Erro ao adicionar anexo: {str(e)}', 'danger')

            return redirect(url_for('contratos.editar', id=id))

        elif acao == 'remover_anexo':
            try:
                anexo_id = request.form.get('anexo_id')
                anexo = ContratoAnexo.query.filter_by(id=int(anexo_id), contrato_id=contrato.id).first_or_404()

                # Remover arquivo do disco
                import os
                if os.path.exists(anexo.caminho_arquivo):
                    os.remove(anexo.caminho_arquivo)

                db.session.delete(anexo)
                db.session.commit()

                flash('Anexo removido!', 'success')
            except Exception as e:
                db.session.rollback()
                flash(f'Erro ao remover anexo: {str(e)}', 'danger')

            return redirect(url_for('contratos.editar', id=id))
        
        elif acao == 'enviar_assinatura':
            try:
                contrato.status = 'aguardando_assinatura'
                db.session.commit()

                _criar_historico(contrato.id, 'assinado', 'Contrato enviado para assinatura')

                flash('Contrato enviado para assinatura!', 'success')
                return redirect(url_for('contratos.index'))

            except Exception as e:
                db.session.rollback()
                flash(f'Erro ao enviar para assinatura: {str(e)}', 'danger')

    # Buscar cláusulas disponíveis para adicionar
    clausulas_disponiveis = ClausulaContratoPadrao.query.filter_by(
        empresa_id=contrato.empresa_id,
        ativo=True
    ).filter(
        ~ClausulaContratoPadrao.id.in_([c.clausula_padrao_id for c in contrato.clausulas if c.clausula_padrao_id])
    ).all()
    
    # Substituir placeholders nas cláusulas
    for clausula in contrato.clausulas:
        clausula.texto_com_placeholders = _substituir_placeholders(clausula.texto, contrato)

    return render_template('editar.html', contrato=contrato, clausulas_disponiveis=clausulas_disponiveis, parametros=parametros)


@contratos_bp.route('/<int:id>/visualizar', methods=['GET'])
@login_required
@require_permission('contratos_view')
def visualizar(id):
    """Visualiza contrato"""
    contrato = scoped_get_or_404(Contrato, id)
    
    # Substituir placeholders nas cláusulas
    for clausula in contrato.clausulas:
        clausula.texto_com_placeholders = _substituir_placeholders(clausula.texto, contrato)
    
    return render_template('visualizar.html', contrato=contrato)


@contratos_bp.route('/<int:id>/assinar', methods=['POST'])
@login_required
@require_permission('contratos_sign')
def assinar(id):
    """Assina contrato"""
    contrato = scoped_get_or_404(Contrato, id)
    
    if contrato.status != 'aguardando_assinatura':
        flash('Apenas contratos aguardando assinatura podem ser assinados.', 'danger')
        return redirect(url_for('contratos.visualizar', id=id))
    
    try:
        contrato.status = 'assinado'
        contrato.data_assinatura = date.today()
        contrato.assinado_por_user_id = current_user.id

        # Criar faturas recorrentes se contrato tiver periodicidade
        if contrato.periodicidade and contrato.periodicidade != 'unico' and contrato.valor_mensal:
            from dateutil.relativedelta import relativedelta

            data_vencimento = contrato.data_inicio_vigencia or date.today()
            numero_faturas = 12  # Padrão: 12 meses

            if contrato.periodicidade == 'trimestral':
                numero_faturas = 4
            elif contrato.periodicidade == 'semestral':
                numero_faturas = 2
            elif contrato.periodicidade == 'anual':
                numero_faturas = 1

            for i in range(numero_faturas):
                if contrato.periodicidade == 'mensal':
                    data_vencimento = data_vencimento + relativedelta(months=1)
                elif contrato.periodicidade == 'trimestral':
                    data_vencimento = data_vencimento + relativedelta(months=3)
                elif contrato.periodicidade == 'semestral':
                    data_vencimento = data_vencimento + relativedelta(months=6)
                elif contrato.periodicidade == 'anual':
                    data_vencimento = data_vencimento + relativedelta(years=1)

                # Criar lançamento financeiro (se modelo existir)
                try:
                    from src.models import Lancamento
                    lancamento = Lancamento(
                        empresa_id=contrato.empresa_id,
                        tipo='receita',
                        descricao=f'Fatura {i+1}/{numero_faturas} - Contrato {contrato.numero}',
                        valor=contrato.valor_mensal,
                        data_vencimento=data_vencimento,
                        data_emissao=date.today(),
                        status='pendente',
                        cliente_id=contrato.cliente_id,
                        contrato_id=contrato.id,
                        criado_por_user_id=current_user.id
                    )
                    db.session.add(lancamento)
                except ImportError:
                    # Modelo Lancamento não existe, apenas log
                    current_app.logger.info(f'Fatura {i+1}/{numero_faturas} seria criada para contrato {contrato.numero}')

        db.session.commit()

        _criar_historico(contrato.id, 'assinado', 'Contrato assinado')

        flash(f'Contrato {contrato.numero} assinado com sucesso!', 'success')
        return redirect(url_for('contratos.visualizar', id=id))
    
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao assinar contrato: {str(e)}', 'danger')
        return redirect(url_for('contratos.visualizar', id=id))


@contratos_bp.route('/<int:id>/cancelar', methods=['POST'])
@login_required
@require_permission('contratos_delete')
def cancelar(id):
    """Cancela contrato"""
    contrato = scoped_get_or_404(Contrato, id)
    
    if contrato.status == 'assinado':
        flash('Não é possível cancelar contratos já assinados.', 'danger')
        return redirect(url_for('contratos.visualizar', id=id))
    
    try:
        motivo = request.form.get('motivo', '')
        contrato.status = 'cancelado'
        contrato.motivo_cancelamento = motivo
        
        db.session.commit()
        
        _criar_historico(contrato.id, 'cancelado', f'Contrato cancelado: {motivo}')
        
        flash(f'Contrato {contrato.numero} cancelado.', 'success')
        return redirect(url_for('contratos.index'))
    
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao cancelar contrato: {str(e)}', 'danger')
        return redirect(url_for('contratos.visualizar', id=id))


@contratos_bp.route('/<int:id>/historico', methods=['GET'])
@login_required
@require_permission('contratos_view')
def historico(id):
    """Visualiza histórico do contrato"""
    contrato = scoped_get_or_404(Contrato, id)
    historico = ContratoHistorico.query.filter_by(contrato_id=id).order_by(ContratoHistorico.versao.desc()).all()
    return render_template('historico.html', contrato=contrato, historico=historico)


@contratos_bp.route('/<int:id>/exportar-pdf', methods=['GET'])
@login_required
@require_permission('contratos_export')
def exportar_pdf(id):
    """Exporta contrato para PDF"""
    from fpdf import FPDF
    import io

    contrato = scoped_get_or_404(Contrato, id)

    if contrato.empresa_id != tenant_id():
        flash('Não autorizado', 'danger')
        return redirect(url_for('contratos.index'))

    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)

        # Título
        pdf.cell(0, 10, 'CONTRATO DE PRESTAÇÃO DE SERVIÇOS', 0, 1, 'C')
        pdf.ln(5)

        # Número
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 8, f'Número: {contrato.numero}', 0, 1, 'C')
        pdf.ln(10)

        # Qualificação das partes
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 8, 'QUALIFICAÇÃO DAS PARTES', 0, 1)
        pdf.ln(5)

        pdf.set_font('Arial', '', 10)
        pdf.cell(0, 6, f'CONTRATADA: {current_user.empresa.nome if current_user.empresa else "-"}', 0, 1)
        pdf.cell(0, 6, f'CNPJ: {current_user.empresa.cnpj if current_user.empresa else "-"}', 0, 1)
        pdf.ln(5)
        pdf.cell(0, 6, f'CONTRATANTE: {contrato.cliente.nome if contrato.cliente else "-"}', 0, 1)
        pdf.cell(0, 6, f'CNPJ/CPF: {contrato.cliente.cnpj_cpf if contrato.cliente else "-"}', 0, 1)
        pdf.ln(10)

        # Cláusulas
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 8, 'CLÁUSULAS', 0, 1)
        pdf.ln(5)

        pdf.set_font('Arial', '', 10)
        for clausula in sorted(contrato.clausulas, key=lambda c: c.ordem):
            pdf.set_font('Arial', 'B', 10)
            pdf.cell(0, 6, f'{clausula.ordem}. {clausula.titulo}', 0, 1)
            pdf.ln(3)

            pdf.set_font('Arial', '', 9)
            # Remover tags HTML e quebrar texto
            texto = clausula.texto.replace('<br>', '\n').replace('<p>', '').replace('</p>', '\n')
            texto = texto.replace('<strong>', '').replace('</strong>', '').replace('<em>', '').replace('</em>', '')
            pdf.multi_cell(0, 5, texto)
            pdf.ln(5)

        # Assinaturas
        pdf.ln(10)
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 8, 'ASSINATURAS', 0, 1)
        pdf.ln(20)
        pdf.cell(80, 8, '_________________', 0, 0, 'C')
        pdf.cell(0, 8, '_________________', 0, 1, 'C')
        pdf.cell(80, 6, 'CONTRATADA', 0, 0, 'C')
        pdf.cell(0, 6, 'CONTRATANTE', 0, 1, 'C')

        # Gerar PDF em memória
        buffer = io.BytesIO()
        pdf_bytes = pdf.output(dest='S').encode('latin-1')
        buffer.write(pdf_bytes)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'contrato_{contrato.numero}.pdf',
            mimetype='application/pdf'
        )

    except Exception as e:
        flash(f'Erro ao exportar PDF: {str(e)}', 'danger')
        return redirect(url_for('contratos.visualizar', id=id))


@contratos_bp.route('/<int:id>/exportar-docx', methods=['GET'])
@login_required
@require_permission('contratos_export')
def exportar_docx(id):
    """Exporta contrato para DOCX"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    contrato = scoped_get_or_404(Contrato, id)

    if contrato.empresa_id != tenant_id():
        flash('Não autorizado', 'danger')
        return redirect(url_for('contratos.index'))

    try:
        doc = Document()

        # Título
        title = doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Número
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f'Número: {contrato.numero}').bold = True

        # Qualificação das partes
        doc.add_heading('QUALIFICAÇÃO DAS PARTES', level=1)
        doc.add_paragraph(f'CONTRATADA: {current_user.empresa.nome if current_user.empresa else "-"}')
        doc.add_paragraph(f'CNPJ: {current_user.empresa.cnpj if current_user.empresa else "-"}')
        doc.add_paragraph('')
        doc.add_paragraph(f'CONTRATANTE: {contrato.cliente.nome if contrato.cliente else "-"}')
        doc.add_paragraph(f'CNPJ/CPF: {contrato.cliente.cnpj_cpf if contrato.cliente else "-"}')

        # Cláusulas
        doc.add_heading('CLÁUSULAS', level=1)
        for clausula in contrato.clausulas.order_by(ContratoClausula.ordem):
            doc.add_heading(f'{clausula.ordem}. {clausula.titulo}', level=2)

            # Remover tags HTML
            texto = clausula.texto_com_placeholders.replace('<br>', '\n').replace('<p>', '').replace('</p>', '\n')
            texto = texto.replace('<strong>', '').replace('</strong>', '').replace('<em>', '').replace('</em>', '')
            texto = texto.replace('<div>', '').replace('</div>', '')
            doc.add_paragraph(texto)

        # Assinaturas
        doc.add_heading('ASSINATURAS', level=1)
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(0, 0)
        cell.text = '_________________'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(0, 1)
        cell.text = '_________________'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(1, 0)
        cell.text = 'CONTRATADA'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(1, 1)
        cell.text = 'CONTRATANTE'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Gerar DOCX em memória
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'contrato_{contrato.numero}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        flash(f'Erro ao exportar DOCX: {str(e)}', 'danger')
        return redirect(url_for('contratos.visualizar', id=id))


@contratos_bp.route('/anexos/<int:id>/baixar', methods=['GET'])
@login_required
@require_permission('contratos_view')
def baixar_anexo(id):
    """Baixa anexo de contrato"""
    anexo = ContratoAnexo.query.get_or_404(id)
    contrato = Contrato.query.get_or_404(anexo.contrato_id)

    if contrato.empresa_id != tenant_id():
        flash('Não autorizado', 'danger')
        return redirect(url_for('contratos.index'))

    try:
        import os
        if os.path.exists(anexo.caminho_arquivo):
            return send_file(
                anexo.caminho_arquivo,
                as_attachment=True,
                download_name=anexo.nome_arquivo,
                mimetype=anexo.tipo_mime
            )
        else:
            flash('Arquivo não encontrado.', 'danger')
            return redirect(url_for('contratos.editar', id=contrato.id))
    except Exception as e:
        flash(f'Erro ao baixar anexo: {str(e)}', 'danger')
        return redirect(url_for('contratos.editar', id=contrato.id))
